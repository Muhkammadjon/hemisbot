import docx
# import asyncio

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return fullText
#print(getText('text.docx'))

def tayyor(file):
    text=[]
    i=1
    
    for bolak in file:
        if bolak=='':
            continue
        if i==1:
            text.extend([bolak, "===="])
        elif i==2:
            x='#'+bolak
            text.extend([x, "===="])
        elif 2<i<=5:
            if i==5:
                text.extend([bolak,"++++"])
            else:
                text.extend([bolak,"===="])
        i+=1
        if i>5:
            i=1
    return text

def yozish(text, filenametax):
    doc = docx.Document()
    for bolak in text:
        doc.add_paragraph(bolak)
    doc.save(filenametax)
    
async def hemis(filename1:str, filename2:str):
    yozish(tayyor(getText(filename1)),filename2)
    
# start_time = time.time()
# yozish(tayyor(getText('text.docx')),'taxladim1.docx')
# end_time = time.time()

# elapsed_time_ms = (end_time - start_time) * 1000

# print(f"Elapsed time: {elapsed_time_ms:.2f} milliseconds")
