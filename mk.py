import docx

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return fullText

def tayyor(file):
    text=[]
    i=1

    for bolak in file:
        if bolak=='':
            continue
        if i==1:
            text.append(bolak)
        elif i==2:
            text.append(bolak)
        elif 2<i<=5:
            if i==5:
                text.extend([bolak,"1"])
            else:
                text.append(bolak)
        i+=1
        if i>5:
            i=1

    return text

def yozish(text, filenametax):
    doc = docx.Document()
    table = doc.add_table(rows=len(text),cols=1, style="Table Grid")
    i=0
    for bolak in text:
        row = table.rows[i].cells
        row[0].text = bolak
        i+=1
    doc.save(filenametax)

async def mk(filename1:str, filename2:str):
    yozish(tayyor(getText(filename1)),filename2)

#yozish(tayyor(getText('text.docx')),'taxladimrow.docx')