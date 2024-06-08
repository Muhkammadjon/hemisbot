import docx

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return fullText
#print(getText('text.docx'))

def tayyor(file):
    text=[]
    i=1
    
    for bolak in file:
        if bolak=='':
            continue
        if i==1:
            text.append(bolak)
        elif i==2:
            x='1. '+bolak
            text.append(x)
        elif 2<i<=5:
            if i==5:
                x='4. '+bolak
                text.extend([x,'Ответ: 1'])
            else:
                x=f'{i-1}. '+bolak
                text.append(x)
        i+=1
        if i>5:
            i=1
    return text

def yozish(text, filenametax):
    doc = docx.Document()
    for bolak in text:
        doc.add_paragraph(bolak)
    doc.save(filenametax)
    
async def modul(filename1:str, filename2:str):
    try :
        yozish(tayyor(getText(filename1)),filename2)
    except:
        print("xato")

# yozish(tayyor(getText('text.docx')),'taxladim_modul.docx')
